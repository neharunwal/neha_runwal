from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
RESUME_DIR = ROOT / "resume"
RESUME_DIR.mkdir(parents=True, exist_ok=True)
DOCX_PATH = RESUME_DIR / "neha-runwal-resume.docx"


def set_cell_shading(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    tc_pr.append(shading)


def set_margins(section, top, right, bottom, left):
    section.top_margin = Inches(top)
    section.right_margin = Inches(right)
    section.bottom_margin = Inches(bottom)
    section.left_margin = Inches(left)


doc = Document()
section = doc.sections[0]
set_margins(section, 0.55, 0.65, 0.55, 0.65)

styles = doc.styles
styles["Normal"].font.name = "Aptos"
styles["Normal"].font.size = Pt(10.3)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = title.add_run("Neha Runwal")
run.font.size = Pt(24)
run.font.bold = True
run.font.color.rgb = RGBColor(20, 31, 57)

subtitle = doc.add_paragraph()
subtitle.space_after = Pt(5)
subtitle_run = subtitle.add_run(
    "Product design, platform engineering, strategic planning, and AI-forward execution"
)
subtitle_run.font.size = Pt(11.2)
subtitle_run.font.color.rgb = RGBColor(62, 86, 126)

meta = doc.add_paragraph()
meta.space_after = Pt(10)
meta_run = meta.add_run(
    "Milpitas, CA | 408-219-7990 | NehaRunwal@gmail.com | linkedin.com/in/neha-runwal-824b4817 | Portfolio created with Codex"
)
meta_run.font.size = Pt(9.2)
meta_run.font.color.rgb = RGBColor(94, 104, 124)

summary_heading = doc.add_paragraph()
summary_heading_run = summary_heading.add_run("Summary")
summary_heading_run.bold = True
summary_heading_run.font.size = Pt(12.3)
summary_heading_run.font.color.rgb = RGBColor(26, 60, 112)

summary = doc.add_paragraph()
summary.add_run(
    "17+ years of extensive experience in product design, strategic planning, development, and execution. "
    "Ready to take on challenges and improve user experience. Recent focus includes learning machine learning, exploring AI tools, and experimenting with vibe coding workflows."
)

table = doc.add_table(rows=1, cols=2)
table.style = "Table Grid"
hdr_cells = table.rows[0].cells
hdr_cells[0].text = "Core Strengths"
hdr_cells[1].text = "Technologies and Platforms"
for cell in hdr_cells:
    set_cell_shading(cell, "D7E8FF")
    for paragraph in cell.paragraphs:
        for row_run in paragraph.runs:
            row_run.bold = True
            row_run.font.color.rgb = RGBColor(20, 31, 57)

row = table.add_row().cells
strengths = [
    "Product design, strategic planning, architecture, development, and execution",
    "Enterprise compliance, security, monitoring, and risk platforms",
    "Team leadership, product strategy, and client-facing requirement translation",
    "Strong mix of startup ownership and large-scale systems delivery",
    "AI-forward learning mindset layered onto mature engineering judgment",
]
techs = [
    "Python, Java, JavaScript, C, C++, Perl, HTML, DHTML, Shell Scripting",
    "Ansible, Ajax, jQuery, UNIX",
    "MySQL, Oracle, Aerospike",
    "NoSQL, Relational Databases, Dashboards, Hybrid Mobile Apps",
    "Machine Learning, AI Tools, Vibe Coding, Codex",
]
for item in strengths:
    row[0].add_paragraph(item, style="List Bullet")
for item in techs:
    row[1].add_paragraph(item, style="List Bullet")

exp_heading = doc.add_paragraph()
exp_heading.space_before = Pt(10)
exp_run = exp_heading.add_run("Professional Experience")
exp_run.bold = True
exp_run.font.size = Pt(12.3)
exp_run.font.color.rgb = RGBColor(26, 60, 112)

experiences = [
    (
        "Stealth Mode Startup",
        "Co-Founder | Apr 2023-Present",
        [
            "Focused on system architecture and product development.",
            "Worked across all product stages including end-to-end design, architecture, and development.",
            "Translated business requirements into technical solutions directly with clients.",
            "Designed and developed the complete tech stack including frontend, mid-tier, and database layers.",
            "Technologies: Python, HTML, JavaScript, MySQL.",
        ],
    ),
    (
        "PayPal, San Jose, USA",
        "Member of Technical Staff | Mar 2012-Mar 2023",
        [
            "Led multiple compliance and security-related projects in the database team.",
            "Developed a monitoring platform for Aerospike, a NoSQL key-value database deployed on 2600+ servers, with future use across other NoSQL and relational databases.",
            "Evaluated roadmap, defined product strategies, designed feature requests, led a team of 6+, and developed solutions based on client requirements.",
            "Previously worked in risk and fraud detection developing applications that tracked health metrics, migrations, and sandbox transitions for 150+ risk services.",
            "Led a team of six to build web applications and a hybrid mobile app for remote health status and alert checks.",
            "Hosted monitoring tech expos and talks, participated in hackathons, encouraged team participation, and won BeatCTF.",
            "Completed machine learning coursework and delivered introductory ML sessions internally for teams.",
            "Technologies: Python, Ansible, Java, Aerospike, MySQL, Oracle.",
        ],
    ),
    (
        "PayPal, San Jose, USA",
        "Summer Intern - Software Engineer | Jun 2011-Nov 2011",
        [
            "Researched and implemented new tools complementing existing monitoring systems in Risk.",
            "Created Python modules interacting with MySQL databases and tables.",
            "Developed a one-pager for application health, alert dashboards, and a web application for critical risk asset visibility.",
            "Showcased three dashboards during the university program and received appreciation from Risk leaders.",
            "Technologies: Python, Ajax, jQuery, Java, MySQL.",
        ],
    ),
    (
        "Infosys, Pune, India",
        "Senior Systems Engineer | Feb 2007-Feb 2010",
        [
            "Designed and developed incoming requirements in C++ using object-oriented programming.",
            "Performed black-box and white-box testing in UNIX.",
            "Automated the testing environment using shell scripting and Perl.",
            "Solved live issues while adhering to standard processes.",
            "Technologies: C, C++, UNIX shell scripting, Perl, DHTML.",
        ],
    ),
    (
        "Cummins College of Engineering, India",
        "Part-time Teaching Assistant | Jul 2006-Dec 2006",
        [
            "Six months of teaching assistantship experience during the early stage of the career.",
        ],
    ),
]

for company, role, bullets in experiences:
    p = doc.add_paragraph()
    p.space_before = Pt(5)
    r = p.add_run(company)
    r.bold = True
    r.font.size = Pt(11)
    p.add_run(f" | {role}")
    for bullet in bullets:
        doc.add_paragraph(bullet, style="List Bullet")

edu_heading = doc.add_paragraph()
edu_heading.space_before = Pt(10)
edu_run = edu_heading.add_run("Education")
edu_run.bold = True
edu_run.font.size = Pt(12.3)
edu_run.font.color.rgb = RGBColor(26, 60, 112)

for bullet in [
    "M.S. in Computer Science, San Jose State University, CA | Dec 2011 | GPA: 3.85",
    "B.E. in Computer Engineering, University of Pune, India | May 2006 | First Class with Distinction",
]:
    doc.add_paragraph(bullet, style="List Bullet")

pub_heading = doc.add_paragraph()
pub_heading.space_before = Pt(10)
pub_run = pub_heading.add_run("Conference Paper and Publication")
pub_run.bold = True
pub_run.font.size = Pt(12.3)
pub_run.font.color.rgb = RGBColor(26, 60, 112)

for bullet in [
    "Presented at SVG Open 2011: “Drag and drop functionality for SVG using jQuery”.",
    "Published online in Journal in Computer Virology on April 3, 2012: “Opcode graph similarity and metamorphic detection”.",
]:
    doc.add_paragraph(bullet, style="List Bullet")

award_heading = doc.add_paragraph()
award_heading.space_before = Pt(10)
award_run = award_heading.add_run("Awards")
award_run.bold = True
award_run.font.size = Pt(12.3)
award_run.font.color.rgb = RGBColor(26, 60, 112)

for bullet in [
    "4 Spot Awards at Infosys for critical initiatives and projects successfully led and delivered.",
    "PayPal Blue Moon Award for successfully delivering Sandbox Risk in 2012 within a few months of joining.",
    "Infosys Bravo Award for developing a weekend support application using Perl and shell script within 3 days.",
]:
    doc.add_paragraph(bullet, style="List Bullet")

doc.save(DOCX_PATH)
print(DOCX_PATH)
